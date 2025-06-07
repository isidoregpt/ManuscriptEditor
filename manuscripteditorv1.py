import streamlit as st
import openai
import anthropic
import os
from io import BytesIO
from docx import Document

# Set up Streamlit
st.set_page_config(page_title="AI Manuscript Editor", layout="wide")
st.title("✍️ AI Manuscript Editor")

# Initialize session state for conversation history
if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = {}
if "feedback_generated" not in st.session_state:
    st.session_state.feedback_generated = False

# Instructions dropdown that remains persistent
with st.expander("📖 How to Use This App & Get API Keys", expanded=False):
    st.markdown("""
    ### 🚀 How to Use This App
    
    1. **Get API Keys** (see links below)
    2. **Enter your API keys** in the fields below
    3. **Select models** you want to use for feedback
    4. **Set an editor persona** (optional - defaults to Nan Graham)
    5. **Input your manuscript** by pasting text or uploading a .txt file
    6. **Write your feedback request** (e.g., "Is the dialogue realistic?", "Does the opening hook the reader?")
    7. **Click "Get Feedback"** to receive AI-powered editorial feedback
    8. **Continue the conversation** with follow-up questions to any model
    9. **Download results** as a Word document for easy sharing and reference
    
    ### 🔑 Where to Get API Keys
    
    **OpenAI API Key** (for GPT models):
    - Sign up at: https://platform.openai.com/
    - Go to API Keys section and create a new key
    - Models available: GPT-4o, GPT-4.1
    
    **Anthropic API Key** (for Claude models):
    - Sign up at: https://console.anthropic.com/
    - Create an API key in your account settings
    - Models available: Claude Sonnet 4
    
    ### 💡 Tips for Best Results
    
    - **Be specific** in your feedback requests (e.g., "Focus on character development in chapter 3")
    - **Provide context** about your genre, target audience, or specific concerns
    - **Use multiple models** to get diverse perspectives on your work
    - **Continue conversations** with follow-up questions for deeper insights
    - **Try different editor personas** to match your writing style or genre needs
    - **Submit manageable chunks** (1-3 pages work best for detailed feedback)
    
    ### 🎯 Example Feedback Requests
    
    - "Does this opening chapter effectively hook the reader?"
    - "Is the dialogue between these characters believable and distinct?"
    - "How is the pacing in this action sequence?"
    - "Are there any plot holes or inconsistencies in this scene?"
    - "Does this character's motivation feel authentic?"
    
    ### 💬 Example Follow-up Questions
    
    - "Can you give me specific examples of how to improve the dialogue?"
    - "What are some alternative ways to start this chapter?"
    - "How would you rewrite this paragraph to improve the pacing?"
    - "Can you elaborate on the character development issues you mentioned?"
    """)

# API key inputs
st.subheader("🔑 API Configuration")
col1, col2 = st.columns(2)

with col1:
    if "openai_api_key" not in st.session_state:
        st.session_state.openai_api_key = ""
    st.session_state.openai_api_key = st.text_input(
        "🤖 Enter your OpenAI API Key",
        value=st.session_state.openai_api_key,
        type="password",
        help="Get your key at: https://platform.openai.com/"
    )

with col2:
    if "anthropic_api_key" not in st.session_state:
        st.session_state.anthropic_api_key = ""
    st.session_state.anthropic_api_key = st.text_input(
        "🧠 Enter your Anthropic API Key",
        value=st.session_state.anthropic_api_key,
        type="password",
        help="Get your key at: https://console.anthropic.com/"
    )

# Set up API clients
openai_client = None
anthropic_client = None

if st.session_state.openai_api_key:
    openai_client = openai.OpenAI(api_key=st.session_state.openai_api_key)

if st.session_state.anthropic_api_key:
    anthropic_client = anthropic.Anthropic(api_key=st.session_state.anthropic_api_key)

# Choose models
openai_models = ["gpt-4o", "gpt-4.1"]
anthropic_models = ["claude-sonnet-4-20250514"]

available_models = []
if openai_client:
    available_models.extend(openai_models)
if anthropic_client:
    available_models.extend(anthropic_models)

if not available_models:
    st.warning("⚠️ Please enter at least one API key to continue.")
    st.info("💡 Don't have API keys? Check the instructions above for sign-up links!")
    st.stop()

st.subheader("⚙️ Model Selection")
selected_models = st.multiselect(
    "🧠 Choose one or more models to use for feedback:",
    options=available_models,
    default=available_models,
    help="Select multiple models to get diverse editorial perspectives on your manuscript."
)

if not selected_models:
    st.warning("Please select at least one model.")
    st.stop()

# Editor configuration
st.subheader("👤 Editor Configuration")
editor_name = st.text_input(
    "Editor Persona", 
    value="Nan Graham",
    help="Name your editor persona (e.g., 'Nan Graham', 'Literary Editor', 'Genre Specialist')"
)

# Manuscript input
st.subheader("📄 Manuscript Input")
manuscript_input = st.text_area(
    "Paste a portion of your manuscript:", 
    height=300,
    help="Paste your text here. For best results, submit 1-3 pages at a time."
)

uploaded_file = st.file_uploader(
    "Or upload a .txt file", 
    type="txt",
    help="Upload a plain text file containing your manuscript excerpt."
)

if uploaded_file:
    manuscript_input = uploaded_file.read().decode("utf-8")
    st.success("✅ Manuscript text loaded from file!")

# Feedback prompt
st.subheader("🎯 Feedback Request")
editor_prompt = st.text_area(
    "What specific feedback would you like?", 
    placeholder="e.g., 'Is the pacing too fast in this chapter?', 'Does the dialogue sound natural?', 'How compelling is this opening?'",
    help="Be specific about what aspects you want feedback on for the most helpful results."
)

# Function to get AI response
def get_ai_response(model, messages, system_prompt):
    try:
        if model in openai_models:
            response = openai_client.chat.completions.create(
                model=model,
                messages=[{"role": "system", "content": system_prompt}] + messages,
                temperature=0.7
            )
            return response.choices[0].message.content.strip()
        elif model in anthropic_models:
            # Convert messages for Anthropic format
            anthropic_messages = []
            for msg in messages:
                if msg["role"] != "system":
                    anthropic_messages.append(msg)
            
            response = anthropic_client.messages.create(
                model=model,
                max_tokens=4096,
                temperature=0.7,
                system=system_prompt,
                messages=anthropic_messages
            )
            return response.content[0].text.strip()
    except Exception as e:
        return f"❌ Error: {str(e)}"

# Run initial feedback
st.subheader("🚀 Generate Feedback")
run_button = st.button("📝 Get Editorial Feedback", type="primary")

if run_button:
    if not manuscript_input or not editor_prompt:
        st.error("❌ Both manuscript text and feedback request are required.")
    else:
        # Reset conversation history for new feedback session
        st.session_state.conversation_history = {}
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        system_prompt = f"You are a brilliant fiction editor named {editor_name}. Provide constructive, detailed feedback on the user's manuscript. Be specific, actionable, and encouraging while identifying areas for improvement."
        initial_message = f"Manuscript:\n{manuscript_input}\n\nFeedback Request:\n{editor_prompt}"
        
        for i, model in enumerate(selected_models):
            status_text.text(f"Getting feedback from {model}...")
            progress_bar.progress((i) / len(selected_models))
            
            # Initialize conversation history for each model
            st.session_state.conversation_history[model] = {
                "messages": [{"role": "user", "content": initial_message}],
                "system_prompt": system_prompt
            }
            
            response = get_ai_response(model, st.session_state.conversation_history[model]["messages"], system_prompt)
            st.session_state.conversation_history[model]["messages"].append({"role": "assistant", "content": response})
        
        progress_bar.progress(1.0)
        status_text.text("✅ Feedback generation complete!")
        st.session_state.feedback_generated = True

# Show results and conversation interface
if st.session_state.feedback_generated and st.session_state.conversation_history:
    st.subheader("📘 Editorial Feedback & Conversation")
    
    # Create tabs for each model
    if len(selected_models) > 1:
        tabs = st.tabs([f"💡 {model}" for model in selected_models])
        
        for i, model in enumerate(selected_models):
            with tabs[i]:
                # Display conversation history
                for j, message in enumerate(st.session_state.conversation_history[model]["messages"]):
                    if message["role"] == "user":
                        if j == 0:  # First message is the original feedback request
                            st.markdown("**📝 Your Original Request:**")
                        else:
                            st.markdown("**❓ Your Follow-up:**")
                        st.markdown(f"*{message['content']}*")
                    elif message["role"] == "assistant":
                        if not message["content"].startswith("❌ Error:"):
                            st.markdown("**🤖 Response:**")
                            st.write(message["content"])
                        else:
                            st.error(message["content"])
                    st.markdown("---")
                
                # Follow-up question input for this model
                follow_up_key = f"follow_up_{model}"
                follow_up = st.text_area(
                    f"💬 Ask a follow-up question to {model}:",
                    key=follow_up_key,
                    placeholder="e.g., 'Can you give specific examples?', 'How would you rewrite this section?'",
                    help="Continue the conversation with this model for deeper insights."
                )
                
                if st.button(f"Send to {model}", key=f"send_{model}"):
                    if follow_up.strip():
                        with st.spinner(f"Getting response from {model}..."):
                            # Add user's follow-up to conversation
                            st.session_state.conversation_history[model]["messages"].append({"role": "user", "content": follow_up})
                            
                            # Get AI response
                            response = get_ai_response(
                                model, 
                                st.session_state.conversation_history[model]["messages"], 
                                st.session_state.conversation_history[model]["system_prompt"]
                            )
                            
                            # Add AI response to conversation
                            st.session_state.conversation_history[model]["messages"].append({"role": "assistant", "content": response})
                        
                        st.rerun()
                    else:
                        st.warning("Please enter a follow-up question.")
    
    else:
        # Single model interface
        model = selected_models[0]
        st.markdown(f"### 💡 Conversation with {model}")
        
        # Display conversation history
        for j, message in enumerate(st.session_state.conversation_history[model]["messages"]):
            if message["role"] == "user":
                if j == 0:  # First message is the original feedback request
                    st.markdown("**📝 Your Original Request:**")
                else:
                    st.markdown("**❓ Your Follow-up:**")
                st.markdown(f"*{message['content']}*")
            elif message["role"] == "assistant":
                if not message["content"].startswith("❌ Error:"):
                    st.markdown("**🤖 Response:**")
                    st.write(message["content"])
                else:
                    st.error(message["content"])
            st.markdown("---")
        
        # Follow-up question input
        follow_up = st.text_area(
            f"💬 Ask a follow-up question to {model}:",
            placeholder="e.g., 'Can you give specific examples?', 'How would you rewrite this section?'",
            help="Continue the conversation with this model for deeper insights."
        )
        
        if st.button(f"Send to {model}"):
            if follow_up.strip():
                with st.spinner(f"Getting response from {model}..."):
                    # Add user's follow-up to conversation
                    st.session_state.conversation_history[model]["messages"].append({"role": "user", "content": follow_up})
                    
                    # Get AI response
                    response = get_ai_response(
                        model, 
                        st.session_state.conversation_history[model]["messages"], 
                        st.session_state.conversation_history[model]["system_prompt"]
                    )
                    
                    # Add AI response to conversation
                    st.session_state.conversation_history[model]["messages"].append({"role": "assistant", "content": response})
                
                st.rerun()
            else:
                st.warning("Please enter a follow-up question.")
    
    # Download functionality
    st.subheader("📥 Export Conversation")
    
    # Generate DOCX file with full conversation
    doc = Document()
    doc.add_heading('AI Manuscript Editor - Full Conversation', 0)
    doc.add_paragraph(f'Editor Persona: {editor_name}')
    doc.add_paragraph(f'Original Feedback Request: {editor_prompt}')
    doc.add_paragraph('')
    doc.add_heading('Original Manuscript Excerpt', level=1)
    doc.add_paragraph(manuscript_input)
    doc.add_paragraph('')
    
    for model in selected_models:
        if model in st.session_state.conversation_history:
            doc.add_heading(f'Conversation with {model}', level=1)
            
            for j, message in enumerate(st.session_state.conversation_history[model]["messages"]):
                if message["role"] == "user":
                    if j == 0:
                        doc.add_heading('Original Request', level=2)
                    else:
                        doc.add_heading('Follow-up Question', level=2)
                    doc.add_paragraph(message["content"])
                elif message["role"] == "assistant" and not message["content"].startswith("❌ Error:"):
                    doc.add_heading('Response', level=2)
                    doc.add_paragraph(message["content"])
                doc.add_paragraph('')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        st.download_button(
            label="📥 Download Full Conversation",
            data=buffer,
            file_name="manuscript_conversation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    with col2:
        if st.button("🔄 Start New Session", type="secondary"):
            st.session_state.conversation_history = {}
            st.session_state.feedback_generated = False
            st.rerun()
    with col3:
        st.success("✨ Continue chatting with your AI editors or download the full conversation!")

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666;'>
        <p>Made with ❤️ for writers everywhere | Powered by OpenAI & Anthropic APIs</p>
    </div>
    """, 
    unsafe_allow_html=True
)
